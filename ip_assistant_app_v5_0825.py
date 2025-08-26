
import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
import smtplib
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email.mime.text import MIMEText
from email.utils import formataddr

# --- Similarity helpers ---
try:
    from rapidfuzz.fuzz import ratio as fuzz_ratio
    def similar(a, b):
        try:
            return fuzz_ratio(a or "", b or "") / 100.0
        except Exception:
            return 0.0
    SIM_ENGINE = "RapidFuzz"
except Exception:
    from difflib import SequenceMatcher
    def similar(a, b):
        try:
            return SequenceMatcher(None, (a or "").lower(), (b or "").lower()).ratio()
        except Exception:
            return 0.0
    SIM_ENGINE = "difflib"

st.set_page_config(page_title="Harshita Legal AI", layout="wide")
#st.title("Harshita Legal AI")
st.markdown(
    """
    <div style='text-align: center;'>
        <div style="font-size:48px; line-height:1;">⚖️</div>
        <h1 style='margin:6px 0 0; color:#2C3E50;'>Harshita Legal AI</h1>
        <p style='font-size:16px; color:#7F8C8D; margin:4px 0 10px;'>
            Renewal reminders • Trademark watch • Contract drafting • Optional email sending
        </p>
        <hr style='margin:0 auto 20px; border: 1px solid #ccc;'/>
    </div>
    """,
    unsafe_allow_html=True
)




with st.expander("About this prototype / How to use"):
    st.write(
        """
        Modules:
        1) Upload IP portfolio -> see renewals due soon -> export reminders CSV or email directly
        2) Upload new trademark filings -> find potential conflicts using fuzzy matching
        3) Generate a Trademark License Agreement (DOCX), optionally convert to PDF
        """
    )
    st.write(f"Similarity engine in use: {SIM_ENGINE}")

st.sidebar.header("Settings")
lead_days = st.sidebar.number_input("Renewal lead time (days)", min_value=1, max_value=365, value=60)
sim_threshold = st.sidebar.slider("Similarity threshold for watch", 0.0, 1.0, 0.75, 0.01)

# ---------- 1) Portfolio & Renewals ----------
st.header("1) Upload IP Portfolio")
portfolio_file = st.file_uploader("Upload ip_portfolio_template.xlsx (or your own with same columns)", type=["xlsx"], key="portfolio")

portfolio = None
due = pd.DataFrame()
if portfolio_file:
    try:
        portfolio = pd.read_excel(portfolio_file, dtype=str).fillna("")
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
        portfolio = None

if portfolio is not None:
    for c in ["FilingDate","RegistrationDate","RenewalDate"]:
        if c in portfolio.columns:
            portfolio[c] = pd.to_datetime(portfolio[c], errors="coerce").dt.date
    st.success(f"Loaded {len(portfolio)} records from portfolio.")
    st.dataframe(portfolio, use_container_width=True)

    st.subheader("Renewals due soon")
    today = pd.to_datetime(datetime.today().date())
    if "RenewalDate" in portfolio.columns:
        dd = portfolio.copy()
        dd["RenewalDate"] = pd.to_datetime(dd["RenewalDate"], errors="coerce")
        dd["DaysToRenewal"] = (dd["RenewalDate"] - today).dt.days
        due = dd[(dd["DaysToRenewal"] >= 0) & (dd["DaysToRenewal"] <= lead_days)].sort_values("DaysToRenewal")
        st.write(f"Found {len(due)} marks due in next {lead_days} days.")
        st.dataframe(due[["Trademark","Class","RegNo","RenewalDate","Owner","OwnerEmail","DaysToRenewal"]], use_container_width=True)

        if len(due):
            msgs = []
            for _, r in due.iterrows():
                msg = (
                    "Subject: Renewal reminder - {tm} (Class {cl}) due {date}\n\n"
                    "Hello {owner},\n\n"
                    "This is a friendly reminder that your trademark \"{tm}\" (Class {cl}, Reg. No. {reg}) is due for renewal on {date}.\n\n"
                    "Please reply to confirm whether you'd like us to proceed with renewal formalities. If yes, we'll share the checklist and fee estimate.\n\n"
                    "Thanks,\nHarshita Legal AI Team"
                ).format(
                    tm=r.get("Trademark",""),
                    cl=r.get("Class",""),
                    date=r.get("RenewalDate",""),
                    owner=r.get("Owner",""),
                    reg=r.get("RegNo","")
                )
                msgs.append({"Owner": r.get("Owner",""), "OwnerEmail": r.get("OwnerEmail",""), "Trademark": r.get("Trademark",""), "Message": msg})
            msgs_df = pd.DataFrame(msgs)
            st.download_button("Download reminder emails (CSV)", data=msgs_df.to_csv(index=False), file_name=f"renewal_reminders_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv", mime="text/csv")

            with st.expander("Optional: Send reminder emails now (SMTP)"):
                smtp_server = st.text_input("SMTP Server (e.g., smtp.office365.com)", value="")
                smtp_port = st.number_input("SMTP Port", value=587, step=1)
                smtp_user = st.text_input("SMTP Username (email address)", value="")
                smtp_pass = st.text_input("SMTP Password / App Password", type="password", value="")
                from_name = st.text_input("From Name", value="Harshita Legal AI Team")
                from_email = st.text_input("From Email", value=smtp_user)

                sel = st.multiselect(
                    "Select marks to email",
                    options=list(due.index),
                    format_func=lambda i: f"{due.loc[i,'Trademark']} (due {due.loc[i,'RenewalDate']})"
                )
                send_click = st.button("Send selected reminders")

                def send_mail(server, port, user, pwd, from_name, from_email, to_email, subject, body):
                    msg = MIMEText(body, "plain")
                    msg["Subject"] = subject
                    msg["From"] = formataddr((from_name, from_email))
                    msg["To"] = to_email
                    with smtplib.SMTP(server, port) as s:
                        s.starttls()
                        s.login(user, pwd)
                        s.sendmail(from_email, [to_email], msg.as_string())

                if send_click:
                    if not (smtp_server and smtp_user and smtp_pass and from_email):
                        st.error("Please fill SMTP settings.")
                    elif not sel:
                        st.warning("Please select at least one mark to email.")
                    else:
                        success, fail = 0, 0
                        for i in sel:
                            r = due.loc[i]
                            body_lines = msgs_df[msgs_df["Trademark"] == r.get("Trademark","")]["Message"]
                            body = body_lines.iloc[0] if len(body_lines) else f"Reminder for {r.get('Trademark','')}"
                            subject = f"Renewal reminder - {r.get('Trademark','')} (Class {r.get('Class','')})"
                            to_email = r.get("OwnerEmail", "")
                            if not to_email:
                                fail += 1
                                continue
                            try:
                                send_mail(smtp_server, smtp_port, smtp_user, smtp_pass, from_name, from_email, to_email, subject, body)
                                success += 1
                            except Exception as e:
                                st.write(f"Error sending to {to_email}: {e}")
                                fail += 1
                        st.success(f"Emails sent: {success}, failed: {fail}")

# ---------- 2) Trademark Watch ----------
st.header("2) Trademark Watch (Prototype)")
filings_file = st.file_uploader("Upload new filings CSV (e.g., new_tm_filings_20rows.csv)", type=["csv"], key="filings")

def base_row_style(row):
    sim = row.get("Similarity", 0)
    if sim >= 0.9:
        return ["background-color: #ffd6d6"] * len(row)  # light red
    elif sim >= 0.8:
        return ["background-color: #ffe8cc"] * len(row)  # light orange
    elif sim >= 0.7:
        return ["background-color: #fff6bf"] * len(row)  # light yellow
    return [""] * len(row)

def highlight_keyword_col(col_series):
    return ['background-color: #fff176; font-weight: bold; border: 1px solid #f1c40f' for _ in col_series]

if (portfolio is not None) and filings_file:
    filings = pd.read_csv(filings_file, dtype=str).fillna("")
    st.success(f"Loaded {len(filings)} new filings.")
    watchwords = []
    if "WatchKeywords" in portfolio.columns:
        for w in portfolio["WatchKeywords"].dropna().astype(str):
            watchwords.extend([x.strip() for x in str(w).split(";") if x.strip()])
    st.write("Watch keywords:", ", ".join(sorted(set(watchwords))[:100]))

    classes_in_portfolio = set(portfolio["Class"].astype(str).tolist()) if "Class" in portfolio.columns else set()

    alerts = []
    for _, f in filings.iterrows():
        mark = f.get("Mark","")
        cl = str(f.get("Class",""))
        if classes_in_portfolio and cl not in classes_in_portfolio:
            continue
        for w in watchwords:
            score = similar(mark, w)
            if score >= sim_threshold:
                alerts.append({
                    "FilingDate": f.get("FilingDate",""),
                    "Mark": mark,
                    "Class": cl,
                    "Applicant": f.get("Applicant",""),
                    "ApplicationNo": f.get("ApplicationNo",""),
                    "MatchedKeyword": w,
                    "Similarity": round(score, 3)
                })

    alerts_df = pd.DataFrame(alerts).drop_duplicates()
    st.subheader("Potential conflicts")

    if len(alerts_df):
        # Filters: per-class and search box
        filter_col1, filter_col2 = st.columns([1,2])
        with filter_col1:
            class_options = sorted(alerts_df["Class"].astype(str).unique().tolist())
            selected_classes = st.multiselect("Filter by class", options=class_options, default=class_options)
        with filter_col2:
            search_text = st.text_input("Search Mark / Applicant / App No.", "").strip().lower()

        filt = alerts_df.copy()
        if selected_classes:
            filt = filt[filt["Class"].astype(str).isin(selected_classes)]
        if search_text:
            mask = (
                filt["Mark"].str.lower().str.contains(search_text) |
                filt["Applicant"].str.lower().str.contains(search_text) |
                filt["ApplicationNo"].str.lower().str.contains(search_text)
            )
            filt = filt[mask]

        filt = filt.sort_values("Similarity", ascending=False)

        # --- Export to Excel: one sheet per class + All_Alerts
        def export_per_class_excel(df):
            bio = BytesIO()
            with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
                # All alerts sheet
                df.to_excel(writer, sheet_name="All_Alerts", index=False)
                # Per-class sheets
                for cl, g in df.groupby(df["Class"].astype(str)):
                    sheet_name = f"Class_{cl}"
                    g.to_excel(writer, sheet_name=sheet_name, index=False)
            bio.seek(0)
            return bio.getvalue()

        excel_bytes = export_per_class_excel(filt)

        # Styling for on-screen table
        styled = filt.style.apply(base_row_style, axis=1).apply(highlight_keyword_col, subset=["MatchedKeyword"])
        st.dataframe(styled, use_container_width=True)

        dl_col1, dl_col2 = st.columns([1,1])
        with dl_col1:
            st.download_button("Download alerts (CSV)", data=filt.to_csv(index=False), file_name=f"tm_watch_alerts_filtered_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv", mime="text/csv")
        with dl_col2:
            st.download_button("Download per-class Excel", data=excel_bytes, file_name=f"tm_watch_alerts_by_class_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        st.caption("Row highlighting by Similarity; MatchedKeyword column highlighted. Filters + search applied to both CSV and Excel exports.")
    else:
        st.info("No potential conflicts found at the current threshold. Try lowering the threshold in the sidebar.")

# ---------- 3) Contract Drafting ----------
st.header("3) Contract Drafting - Trademark License")

with st.form("license_form_v2"):
    col1, col2, col3 = st.columns(3)
    with col1:
        licensor = st.text_input("LicensorName", "Acme Foods Pvt Ltd")
        licensor_addr = st.text_input("LicensorAddress", "Mumbai, India")
        licensee = st.text_input("LicenseeName", "SnackCo Ltd")
        licensee_addr = st.text_input("LicenseeAddress", "New Delhi, India")
    with col2:
        trademark = st.text_input("Trademark", "BrandX")
        clss = st.text_input("Class", "30")
        territory = st.text_input("Territory", "India")
    with col3:
        lic_type = st.selectbox("LicenseType", ["exclusive","non-exclusive","sole"])
        eff = st.text_input("EffectiveDate", datetime.today().date().isoformat())
        term = st.text_input("TermYears", "3")
        royalty = st.text_input("RoyaltyPercent", "5")
        law = st.text_input("GoverningLaw", "Laws of India")
        seat = st.text_input("ArbitrationSeat", "New Delhi")

    submitted = st.form_submit_button("Generate Agreement (.docx)")

if submitted:
    template_candidates = [
        "TM_License_Template_Placeholders.docx",
        "./TM_License_Template_Placeholders.docx",
        "/mnt/data/TM_License_Template_Placeholders.docx",
    ]
    doc = None
    for p in template_candidates:
        try:
            doc = Document(p)
            break
        except Exception:
            continue
    if doc is None:
        # Title
        title = doc.add_heading("TRADEMARK LICENSE AGREEMENT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.size = Pt(16)          # bigger than normal (default is ~11 or 12)
        run.bold = True

# Add a small spacing after title
        doc.add_paragraph("")
 # --- Body content ---       
        doc.add_paragraph(
        "This Trademark License Agreement (the 'Agreement') is made between "
        "{{LicensorName}} ('Licensor') and {{LicenseeName}} ('Licensee') "
        "for the trademark {{Trademark}}.")

# --- Ensure all in one page ---
# Reduce page margins
    for section in doc.sections:
        section.top_margin = Pt(36)     # ~0.5 inch
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)    # ~0.75 inch
        section.right_margin = Pt(54)

# --- Add footer ---
    #   section = doc.sections[0]                  # get section
     #   footer = section.footer                     # access footer
    # Add a new paragraph in the footer (or overwrite the first one)
      #  footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

       # footer_para.text = "⚖️ Shared by Harshita Legal Team – Confidential"
       # footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Ensure footer appears on every page (all sections)
    for section in doc.sections:
    # If the template has different first-page footer, turn it off OR write to both
        try:
            section.different_first_page_header_footer = False
        except Exception:
            pass

    # Main footer (all pages)
        ftr = section.footer
        para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        para.text = "⚖️ Shared by IP Team – Confidential"
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Also set first-page footer just in case the flag is still respected
        first = section.first_page_footer
        p1 = first.paragraphs[0] if first.paragraphs else first.add_paragraph()
        p1.text = "⚖️ Shared by Harshita Legal AI Team – Confidential"
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER


    repl = {
        "{{LicensorName}}": licensor,
        "{{LicensorAddress}}": licensor_addr,
        "{{LicenseeName}}": licensee,
        "{{LicenseeAddress}}": licensee_addr,
        "{{Trademark}}": trademark,
        "{{Class}}": clss,
        "{{Territory}}": territory,
        "{{LicenseType}}": lic_type,
        "{{EffectiveDate}}": eff,
        "{{TermYears}}": term,
        "{{RoyaltyPercent}}": royalty,
        "{{GoverningLaw}}": law,
        "{{ArbitrationSeat}}": seat,
    }

    for p in doc.paragraphs:
        for k, v in repl.items():
            if k in p.text:
                for run in p.runs:
                    run.text = run.text.replace(k, str(v))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in repl.items():
                    if k in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = run.text.replace(k, str(v))

    bio = BytesIO()
    doc.save(bio)
    st.session_state["license_doc"] = bio.getvalue()

if "license_doc" in st.session_state:
    st.download_button(
        "Download License Agreement",
        data=st.session_state["license_doc"],
        file_name=f"Trademark_License_Agreement__{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

st.divider()
st.caption("Prototype only. Use official IP data sources & legal review before action. MatchedKeyword cells highlighted; filters + search; export to per-class Excel.")
